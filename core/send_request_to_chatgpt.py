from django.db.models import Max, Min
from django.core.files.base import ContentFile

from .models import Paragraphs, Videos, Screenshots, Articles

from profbuh_video_to_article.config import *

import time
import openai
from openai import OpenAIError
import docx
import tempfile
from docx import Document


class ArticleCreation:
    def __init__(self, video_id, article_legth, annotation_length):
        self.article_legth = article_legth
        self.annotation_length = annotation_length
        self.video_id = video_id
        self.article_total_length = 0
        openai.api_key = OPEN_AI_KEY

    def get_response(self):
        paragraphs = Paragraphs.objects.filter(video=self.video_id)
        doc = Document()
        video = Videos.objects.get(id=self.video_id)
        title = video.video_name
        doc.add_heading(title, level=1)
        blank_paragraph = doc.add_paragraph()

        # Вычисление общей длины статьи
        if self.article_legth != "Безгранично":
            for paragraph in paragraphs:
                self.article_total_length += paragraph.len_article

        # Удаление первой группы скриншотов
        screenshots_to_delete = Screenshots.objects.filter(group=1)
        for screenshot in screenshots_to_delete:
            screenshot.screenshot.delete()  # Удаление связанного медиафайла
        screenshots_to_delete.delete()

        # Удаление последней группы скриншотов
        last_group = Screenshots.objects.values(
            'group').order_by('-group').first()
        if last_group:
            screenshots_to_delete = Screenshots.objects.filter(
                group=last_group['group'])
            for screenshot in screenshots_to_delete:
                screenshot.screenshot.delete()  # Удаление связанного медиафайла
            screenshots_to_delete.delete()

        full_text = []
        for paragraph in paragraphs:
            if self.article_legth != "Безгранично":
                proportion = (paragraph.article /
                              self.article_total_length)*self.article_legth
            else:
                proportion = "Безгранично"

            # Получение временных меток групп
            group_timestamps = (
                Screenshots.objects.values('group')
                .annotate(min_timestamp=Min('timestamp'), max_timestamp=Max('timestamp'))
            )

            # Перебор временных меток групп и удаление групп с разницей < 10
            for group_data in group_timestamps:
                group = group_data['group']
                min_timestamp = group_data['min_timestamp']
                max_timestamp = group_data['max_timestamp']
                if max_timestamp - min_timestamp < 10:
                    # Удаление скриншотов и связанных медиафайлов для группы
                    screenshots_to_delete = Screenshots.objects.filter(
                        group=group)
                    for screenshot in screenshots_to_delete:
                        screenshot.screenshot.delete()  # Удаление связанного медиафайла
                    screenshots_to_delete.delete()  # Удаление скриншотов из таблицы

            prompt = """
            Задание: Перефразируйте текст таким образом, чтобы сохранить его смысл.
            Тема статьи: {}
            Стиль: Официальный. 
            Запрещено: придумывать или добавлять что-то новое, использовать прямую речь, обращаться к зрителям или читателям (добрый день, уважаемые зрители), использовать слово "мы", указывать тему статьи, писать "В статье обсуждается/рассматривается".
            Разрешено: использовать списки.
            Формат ответа: текст.
            Максимальное количество символов: {}

            Текст: {}
            """.format(title, proportion, paragraph.article)
            
            max_retries = 5
            retries = 0
            # Отправка запроса на перефразировку модели ChatGPT
            retries = 0  # Количество повторных попыток
            max_retries = 3  # Максимальное количество повторных попыток

            while retries < max_retries:
                try:
                    # Отправка запроса на перефразировку модели ChatGPT
                    response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=[
                        {"role": "system", "content": "You are a helpful assistant."},
                        {"role": "user", "content": prompt}
                    ])
                    break  # Прервать цикл, если запрос успешен
                except OpenAIError as e:
                    # Обработка ошибки OpenAIError (например, запись ошибки в журнал, учет ограничений скорости)
                    print(f"Ошибка: {e}")
                    retries += 1
                    if retries < max_retries:
                        print(f"Повторная попытка через {60} секунд...")
                        time.sleep(60)
            else:
                print("Превышено максимальное количество попыток. Запрос не удался.")
                        

            # Извлечение сгенерированной перефразировки из ответа
            paraphrase = response.choices[0]['message']['content']
            decoded_content = paraphrase.encode('utf-8').decode()

            nearest_group = Screenshots.objects.filter(
                timestamp__gte=paragraph.timestamp).order_by('timestamp').values('group').first()
            nearest_group = nearest_group['group']
            print(nearest_group)

            if nearest_group:
                # Получение первого, среднего и последнего скриншотов из группы
                screenshots = Screenshots.objects.filter(
                    group=nearest_group).order_by('timestamp')
                first_screenshot = screenshots.first()
                last_screenshot = screenshots.last()
                middle_screenshot = screenshots[(screenshots.count() - 1) // 2]
                screenshots_to_delete = screenshots.exclude(
                    id__in=[first_screenshot.id, middle_screenshot.id, last_screenshot.id])
                for screenshot in screenshots_to_delete:
                    screenshot.screenshot.delete()  # Удаление связанного медиафайла
                screenshots_to_delete.delete()

            hours = paragraph.timestamp // 3600  # Получение целочисленного деления для часов
            # Получение целочисленного деления для минут
            minutes = (paragraph.timestamp % 3600) // 60
            seconds = paragraph.timestamp % 60  # Получение остатка от деления для секунд

            timestamp_formatted = "{:02d}:{:02d}:{:02d}".format(
                hours, minutes, seconds)
            doc.add_paragraph(
                f"Время начала абзаца: {timestamp_formatted}", style="BodyText")
            full_text.append(decoded_content)
            doc.add_paragraph(decoded_content, style="BodyText")
            doc.add_picture(first_screenshot.screenshot.path, width=docx.shared.Inches(
                6), height=docx.shared.Inches(4))
            doc.add_picture(middle_screenshot.screenshot.path, width=docx.shared.Inches(
                6), height=docx.shared.Inches(4))
            doc.add_picture(last_screenshot.screenshot.path, width=docx.shared.Inches(
                6), height=docx.shared.Inches(4))
            print('One more iteration done')
            time.sleep(30)

        prompt = """
            Напиши аннотацию к статье
            Максимальное количество символов: {}

            Текст: {}
            """.format(self.annotation_length, full_text[0])

        # Отправка запроса наполучение аннотации от модели ChatGPT
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ]
        )

        # Извлечение сгенерированной аннотации из ответа
        annotation = response.choices[0]['message']['content']
        decoded_content = annotation.encode('utf-8').decode()

        blank_paragraph.text = f"Аннотация: {decoded_content}"

        # Сохранение документа во временном файле
        with tempfile.NamedTemporaryFile(suffix='.docx') as temp_file:
            doc.save(temp_file.name)

            # Чтение содержимого временного файла
            with open(temp_file.name, 'rb') as file:
                file_content = file.read()

        # Создание ContentFile из содержимого файла
        article_file_content = ContentFile(file_content, name='output.docx')

        # Создание нового экземпляра Articles
        video_instance = Videos.objects.get(id=self.video_id)
        article_instance = Articles(video=video_instance, article_id=self.video_id)

        # Присвоение файла статьи article_instance
        article_instance.article_file.save('output.docx', article_file_content)
        article_instance.save()

